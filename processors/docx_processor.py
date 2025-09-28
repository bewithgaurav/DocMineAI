"""
DOCX document processor using python-docx for Word document extraction
"""
from pathlib import Path
from typing import Dict, Any, List
import logging

try:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph
except ImportError:
    Document = None
    logging.warning("python-docx not installed. Run: pip install python-docx")

from .base_processor import BaseProcessor

class DocxProcessor(BaseProcessor):
    """Processor for DOCX documents with structure preservation"""
    
    def __init__(self, config: Dict[str, Any] = None):
        super().__init__(config)
        
        if Document is None:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
    
    def can_process(self, file_path: Path) -> bool:
        """Check if file is a DOCX"""
        return file_path.suffix.lower() in ['.docx', '.doc']
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from DOCX preserving structure"""
        try:
            doc = Document(file_path)
            text_content = []
            
            # Process document elements in order
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # Paragraph
                    paragraph = Paragraph(element, doc)
                    if paragraph.text.strip():
                        # Check if it's a heading style
                        if paragraph.style.name.startswith('Heading'):
                            level = paragraph.style.name.replace('Heading ', '')
                            text_content.append(f"{'#' * int(level if level.isdigit() else 1)} {paragraph.text}")
                        else:
                            text_content.append(paragraph.text)
                
                elif isinstance(element, CT_Tbl):
                    # Table
                    table = Table(element, doc)
                    text_content.append("--- Table ---")
                    
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip().replace('\n', ' ')
                            row_text.append(cell_text)
                        if any(row_text):  # Only add non-empty rows
                            text_content.append(" | ".join(row_text))
                    
                    text_content.append("--- End Table ---")
            
            full_text = "\n".join(text_content)
            return self.clean_text(full_text)
            
        except Exception as e:
            self.logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            return ""
    
    def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from DOCX"""
        metadata = self.get_file_info(file_path)
        
        try:
            doc = Document(file_path)
            
            # Core properties
            core_props = doc.core_properties
            metadata.update({
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "category": core_props.category or "",
                "comments": core_props.comments or "",
                "keywords": core_props.keywords or "",
                "language": core_props.language or "",
                "last_modified_by": core_props.last_modified_by or "",
                "created": core_props.created.isoformat() if core_props.created else "",
                "modified": core_props.modified.isoformat() if core_props.modified else "",
            })
            
            # Document structure analysis
            paragraphs = len(doc.paragraphs)
            tables = len(doc.tables)
            
            # Count headings
            heading_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    heading_count += 1
            
            # Count images (embedded objects)
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
            
            metadata.update({
                "paragraphs": paragraphs,
                "tables": tables,
                "headings": heading_count,
                "images": image_count,
            })
            
        except Exception as e:
            self.logger.error(f"Error extracting DOCX metadata from {file_path}: {e}")
            metadata["docx_error"] = str(e)
        
        return metadata
    
    def extract_tables(self, file_path: Path) -> List[List[List[str]]]:
        """Extract tables from DOCX as structured data"""
        tables = []
        
        try:
            doc = Document(file_path)
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    table_data.append(row_data)
                tables.append(table_data)
                
        except Exception as e:
            self.logger.error(f"Error extracting tables from DOCX {file_path}: {e}")
        
        return tables
    
    def extract_headings(self, file_path: Path) -> List[Dict[str, Any]]:
        """Extract document structure based on headings"""
        headings = []
        
        try:
            doc = Document(file_path)
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name.replace('Heading ', '')
                    headings.append({
                        "text": paragraph.text,
                        "level": int(level) if level.isdigit() else 1,
                        "style": paragraph.style.name,
                        "position": i
                    })
                    
        except Exception as e:
            self.logger.error(f"Error extracting headings from DOCX {file_path}: {e}")
        
        return headings